from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parents[1] / "public"
DOCX_PATH = OUT_DIR / "Medvedskiy_Pavlo_Resume.docx"

INK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
CYAN = RGBColor(8, 145, 178)
AMBER = RGBColor(146, 64, 14)
BORDER = "CBD5E1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bottom_border(paragraph, color=BORDER, size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_run(run, size=None, bold=False, color=None, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_text(paragraph, text, size=9, bold=False, color=INK, italic=False):
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return run


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    add_text(p, text.upper(), size=10, bold=True, color=CYAN)
    add_bottom_border(p, color="67E8F9", size="6")
    return p


def add_bullet(doc, text, level=0, size=8.7):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22 + level * 0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.7)
    p.paragraph_format.line_spacing = 1.03
    add_text(p, text, size=size, color=INK)
    return p


def add_role(doc, company, role, period, context=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    add_text(p, company, size=9.3, bold=True, color=INK)
    add_text(p, f" - {role}", size=9.3, color=INK)
    add_text(p, f" | {period}", size=9.1, bold=True, color=AMBER)
    if context:
        c = doc.add_paragraph()
        c.paragraph_format.space_after = Pt(1.5)
        c.paragraph_format.line_spacing = 1.03
        add_text(c, context, size=8.5, color=MUTED, italic=True)


def add_skill_line(doc, label, items):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.8)
    p.paragraph_format.line_spacing = 1.03
    add_text(p, f"{label}: ", size=8.8, bold=True, color=INK)
    add_text(p, ", ".join(items), size=8.8, color=INK)


def add_two_col_section(doc, left_title, left_items, right_title, right_items):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.45)
    table.columns[1].width = Inches(3.45)
    for cell in table.rows[0].cells:
        cell.width = Inches(3.45)
        set_cell_shading(cell, "F8FAFC")
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
    for cell, title, items in [
        (table.rows[0].cells[0], left_title, left_items),
        (table.rows[0].cells[1], right_title, right_items),
    ]:
        p = cell.paragraphs[0]
        add_text(p, title, size=8.6, bold=True, color=CYAN)
        for item in items:
            bp = cell.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.18)
            bp.paragraph_format.first_line_indent = Inches(-0.1)
            bp.paragraph_format.space_after = Pt(1.2)
            add_text(bp, item, size=8.3, color=INK)


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9)
    normal.font.color.rgb = INK

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    add_text(p, "PAVLO MEDVEDSKYI", size=19, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    add_text(p, "Senior Quality Assurance Engineer", size=11.5, bold=True, color=CYAN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    add_text(
        p,
        "Kharkiv, Ukraine | Medvedskiypa@gmail.com | t.me/Wisll | "
        "linkedin.com/in/pavlo-medvedskyi-74231913b | pavlo-medvedskyi.github.io",
        size=8.2,
        color=MUTED,
    )
    add_bottom_border(p)

    add_heading(doc, "Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    add_text(
        p,
        "Senior Quality Assurance Engineer with 8+ years of commercial experience in desktop, web, API, "
        "database, and real-time systems testing. Experienced in functional, integration, regression, "
        "exploratory, smoke, release, and system testing. Strong background in requirements analysis, "
        "test design, SQL, API validation, databases, network protocols, Electron applications, Docker, "
        "Linux, and cross-platform validation. Comfortable investigating defects through logs, data state, "
        "client-server behavior, and business logic, with mentoring and knowledge-sharing experience.",
        size=8.9,
        color=INK,
    )

    add_heading(doc, "Core Technical Skills")
    add_skill_line(
        doc,
        "Testing",
        [
            "Functional",
            "Integration",
            "Regression",
            "Smoke",
            "Exploratory",
            "System",
            "Release Validation",
            "Cross-platform Testing",
        ],
    )
    add_skill_line(
        doc,
        "Technical QA",
        [
            "Requirements Analysis",
            "Test Design",
            "API Testing",
            "SQL",
            "Database Validation",
            "Log Analysis",
            "Defect Investigation",
            "Client-server Validation",
        ],
    )
    add_skill_line(
        doc,
        "Platforms",
        [
            "Desktop Applications",
            "Web Applications",
            "Mobile",
            "Electron",
            "Ubuntu",
            "Windows",
            "macOS",
            "Linux",
            "Docker",
        ],
    )
    add_skill_line(
        doc,
        "Tools",
        [
            "Postman",
            "Swagger",
            "Jira",
            "Zephyr",
            "TestRail",
            "Confluence",
            "Fiddler",
            "Chrome DevTools",
            "Sentry",
            "Google Cloud Console",
            "OpenShift",
            "VirtualBox",
        ],
    )
    add_skill_line(
        doc,
        "Automation Exposure",
        ["Java", "Rest Assured", "TestNG", "Postman runners", "Selenide/Selenium basics"],
    )
    add_skill_line(
        doc,
        "Databases",
        ["PostgreSQL", "MSSQL", "Derby DB", "MongoDB", "DynamoDB", "BigQuery", "DBeaver"],
    )

    add_heading(doc, "Work Experience")
    add_role(
        doc,
        "CODY SOLUTIONS",
        "Senior QA Engineer",
        "Jan 2025 - Present",
        "Confidential NDA desktop platform - large-scale Electron application with cross-platform delivery.",
    )
    for item in [
        "Currently working on a technically complex Electron desktop application under NDA, with validation across Ubuntu, Windows, and macOS.",
        "Own functional, integration, regression, exploratory, system, and release validation activities for desktop and connected backend functionality.",
        "Analyze requirements, clarify edge cases, design test coverage, and validate business logic across client-server workflows.",
        "Investigate defects using logs, API responses, database state, simulator behavior, and environment-specific reproduction steps.",
        "Perform API and database validation, including SQL checks and integration verification across backend and desktop layers.",
        "Support QA documentation, test reporting, regression planning, onboarding, mentoring, and knowledge sharing within Agile delivery.",
    ]:
        add_bullet(doc, item)

    add_role(
        doc,
        "1648 FACTORY",
        "Senior QA Engineer",
        "2024",
        "Fintech web application for cryptocurrency trading.",
    )
    for item in [
        "Performed manual, API, and database testing across multiple environments, including PostgreSQL validation.",
        "Created and executed automated Postman collections and runners for API regression coverage.",
        "Validated MetaMask integration, transaction flows, and key web application scenarios.",
        "Executed smoke and regression testing for web builds; maintained test cases and QA documentation.",
        "Collaborated with developers and product owners to clarify requirements and support Scrum delivery.",
    ]:
        add_bullet(doc, item)

    add_role(
        doc,
        "RAILSWARE",
        "QA Engineer",
        "2022 - 2023",
        "Coupler.io - SaaS data synchronization platform.",
    )
    for item in [
        "Validated data synchronization workflows, API integrations, and database behavior for scheduled data transfers.",
        "Tested integrations with Facebook Ads, LinkedIn Ads, Google Ads, Shopify, WooCommerce, and other external services.",
        "Performed API, database, regression, smoke, responsive, and mobile testing for weekly releases.",
        "Analyzed production and staging issues using Google Cloud Console, Sentry, logs, and product support context.",
        "Maintained test plans, checklists, documentation, and cross-functional communication with product and engineering teams.",
    ]:
        add_bullet(doc, item)

    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    add_heading(doc, "Work Experience Continued")
    add_role(
        doc,
        "SOFTSERVE",
        "QA Engineer",
        "2019 - 2022",
        "Service delivery and monitoring solutions across VoIP, SIP, and MPLS network environments.",
    )
    for item in [
        "Analyzed functional requirements and created QA strategy, test cases, regression suites, and smoke coverage.",
        "Performed web, API, installation, and compatibility testing on Windows and Unix/CentOS environments.",
        "Validated backend data integrity with PostgreSQL and tested API behavior using Postman and Fiddler.",
        "Investigated defects through application/server logs, OpenShift environments, Java applications, and deployment context.",
        "Worked with VirtualBox, Chocolatey, Google Cloud, OpenShift, and distributed client-server environments.",
        "Mentored junior QA engineers, reviewed test cases, supported onboarding, and participated in QA appraisals.",
    ]:
        add_bullet(doc, item)

    add_role(
        doc,
        "EPAM SYSTEMS",
        "QA Engineer",
        "2018 - 2019",
        "EPAM Cloud - multi-cloud orchestrator for AWS, GCP, and Azure.",
    )
    for item in [
        "Performed manual functional, UI, database, CLI, and mobile testing for cloud management workflows.",
        "Created test cases, documented defects and test results in Jira, and supported Agile QA activities.",
        "Validated data behavior in MongoDB and DynamoDB and tested cloud-provider related scenarios.",
        "Collaborated with developers and QA team members during daily standups, reviews, and regression cycles.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Selected Project Exposure")
    add_two_col_section(
        doc,
        "Confidential Enterprise Desktop Platform",
        [
            "Electron desktop QA",
            "Ubuntu / Windows / macOS",
            "Simulator testing",
            "Release validation",
            "API, DB, logs, and integration checks",
        ],
        "SaaS / Fintech / Data Platforms",
        [
            "Web, mobile, API, and database testing",
            "External API integrations and ETL flows",
            "Crypto transaction and wallet-related workflows",
            "Regression, smoke, exploratory, and support-driven testing",
            "Requirements analysis and documentation",
        ],
    )

    add_heading(doc, "Professional Focus")
    for item in [
        "Technically challenging products where quality depends on understanding system behavior, integrations, and data flow.",
        "Strong preference for evidence-based defect investigation using logs, API responses, database state, and reproducible steps.",
        "Able to work independently across the full SDLC while collaborating closely with product, development, and support teams.",
    ]:
        add_bullet(doc, item)

    # Quiet footer.
    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_before = Pt(0)
        add_text(footer, "Pavlo Medvedskyi | Senior Quality Assurance Engineer", size=7.5, color=MUTED)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
